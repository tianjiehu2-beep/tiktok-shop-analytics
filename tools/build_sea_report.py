"""Build SEA hot-category analysis workbook (xlsx) + selection report (docx).

Data comes from data/tiktok_shop.db rows collected by tools/collect_socialcrawl.py
(real TikTok Shop TH data via SocialCrawl API).

Usage:
  python tools/build_sea_report.py [--db data/tiktok_shop.db] [--outdir reports]
"""

from __future__ import annotations

import argparse
import math
import sqlite3
import sys
from datetime import date, datetime, timezone
from pathlib import Path

CATEGORIES = [
    "美妆个护-护肤", "美妆个护-彩妆", "服饰-女装", "服饰-男装",
    "3C数码-手机配件", "家居生活-家居装饰", "食品饮料-零食", "健康保健-保健品",
    "宠物用品", "运动户外-健身器材", "时尚饰品-珠宝", "母婴用品",
]


def utc_today() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d")


def load_products(db_path: str) -> list[dict]:
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    placeholders = ",".join("?" for _ in CATEGORIES)
    rows = conn.execute(
        f"""SELECT product_id, title, category, price, original_price, sold_count,
                   rating, seller_name, seller_id
            FROM products
            WHERE category IN ({placeholders}) AND is_active = 1
            ORDER BY sold_count DESC""",
        CATEGORIES,
    ).fetchall()
    conn.close()
    products = []
    for r in rows:
        p = dict(r)
        p["discount"] = (1 - p["price"] / p["original_price"]) if (p["original_price"] and p["original_price"] > p["price"]) else 0.0
        p["gmv"] = (p["price"] or 0.0) * (p["sold_count"] or 0)
        products.append(p)
    return products


def category_summary(products: list[dict]) -> list[dict]:
    rows = []
    for cat in CATEGORIES:
        items = [p for p in products if p["category"] == cat]
        if not items:
            continue
        prices = sorted(p["price"] for p in items if p["price"] > 0)
        med = prices[len(prices) // 2] if prices else 0.0
        avg = sum(p["price"] for p in items) / len(items)
        rows.append({
            "category": cat,
            "count": len(items),
            "avg_price": avg,
            "median_price": med,
            "max_sold": max(p["sold_count"] or 0 for p in items),
            "total_sold": sum(p["sold_count"] or 0 for p in items),
            "avg_rating": round(
                sum(r for p in items if (r := p["rating"] or 0) > 0)
                / max(sum(1 for p in items if (p["rating"] or 0) > 0), 1), 2),
            "avg_discount": sum(p["discount"] for p in items) / len(items),
            "total_gmv": sum(p["gmv"] for p in items),
        })
    rows.sort(key=lambda r: r["total_sold"], reverse=True)
    return rows


def blue_ocean_score(p: dict, shop_freq: dict, max_freq: int) -> float:
    demand = min(math.log10(max(p["sold_count"] or 1, 1)) / 7, 1.0)
    quality = (p["rating"] or 0.0) / 5.0
    competition = 1.0 - (shop_freq.get(p["seller_id"], 0) / max(max_freq, 1))
    return 0.4 * demand + 0.3 * quality + 0.3 * competition


def build_xlsx(products: list[dict], summaries: list[dict], out: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    head_fill = PatternFill("solid", fgColor="1F4E79")
    head_font = Font(bold=True, color="FFFFFF")

    # Sheet 1: category overview
    ws = wb.active
    ws.title = "类目总览"
    headers = ["类目", "商品数", "平均售价(THB)", "中位价(THB)", "最高销量", "总销量",
               "平均评分", "平均折扣率", "预估GMV(THB)"]
    ws.append(headers)
    for c in ws[1]:
        c.fill, c.font = head_fill, head_font
    for r in summaries:
        ws.append([r["category"], r["count"], round(r["avg_price"], 2), round(r["median_price"], 2),
                   r["max_sold"], r["total_sold"], r["avg_rating"],
                   f'{r["avg_discount"]:.0%}', round(r["total_gmv"])])
    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 16
    ws.freeze_panes = "A2"

    # Sheet 2: top products
    ws2 = wb.create_sheet("热销商品Top100")
    h2 = ["排名", "标题", "类目", "售价(THB)", "原价(THB)", "折扣率", "累计销量", "评分", "店铺", "预估GMV(THB)"]
    ws2.append(h2)
    for c in ws2[1]:
        c.fill, c.font = head_fill, head_font
    top = sorted(products, key=lambda p: p["sold_count"] or 0, reverse=True)[:100]
    for i, p in enumerate(top, 1):
        ws2.append([i, p["title"], p["category"], round(p["price"] or 0, 2),
                    round(p["original_price"] or 0, 2), f'{p["discount"]:.0%}',
                    p["sold_count"] or 0, p["rating"] or 0, p["seller_name"],
                    round(p["gmv"])])
    widths2 = [6, 60, 18, 12, 12, 9, 12, 8, 22, 16]
    for i, w in enumerate(widths2, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = "A2"

    # Sheet 3: blue ocean opportunities
    shop_freq: dict = {}
    for p in products:
        shop_freq[p["seller_id"]] = shop_freq.get(p["seller_id"], 0) + 1
    max_freq = max(shop_freq.values()) if shop_freq else 1
    scored = []
    for p in products:
        if (p["sold_count"] or 0) < 100 or (p["sold_count"] or 0) > 300000:
            continue
        scored.append((blue_ocean_score(p, shop_freq, max_freq), p))
    scored.sort(key=lambda x: x[0], reverse=True)
    ws3 = wb.create_sheet("蓝海机会Top50")
    h3 = ["排名", "蓝海分", "标题", "类目", "售价(THB)", "累计销量", "评分", "店铺"]
    ws3.append(h3)
    for c in ws3[1]:
        c.fill, c.font = head_fill, head_font
    for i, (score, p) in enumerate(scored[:50], 1):
        ws3.append([i, round(score, 3), p["title"], p["category"],
                    round(p["price"] or 0, 2), p["sold_count"] or 0,
                    p["rating"] or 0, p["seller_name"]])
    widths3 = [6, 9, 60, 18, 12, 12, 8, 22]
    for i, w in enumerate(widths3, 1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.freeze_panes = "A2"

    # Sheet 4: notes
    ws4 = wb.create_sheet("数据说明")
    notes = [
        ["数据源", "SocialCrawl API（TikTok Shop 泰国站实时搜索数据）"],
        ["采集日期", utc_today()],
        ["采集方式", "12 个东南亚热门类目关键词，每类目 Top30"],
        ["样本量", f"{len(products)} 条商品记录"],
        ["币种", "THB（泰铢），预估GMV = 累计销量 × 当前售价"],
        ["说明", "销量为商品累计销量；折扣率 = 1 - 现价/原价"],
        ["说明2", "蓝海分 = 0.4×需求 + 0.3×质量 + 0.3×低竞争（销量100~30万区间内）"],
        ["注意", "部分类目存在 1 泰铢引流款/0 评分新品，分析时需结合评分与店铺资质"],
    ]
    for row in notes:
        ws4.append(row)
    ws4.column_dimensions["A"].width = 12
    ws4.column_dimensions["B"].width = 80

    wb.save(out)
    print(f"xlsx saved: {out}")


def build_docx(products: list[dict], summaries: list[dict], out: Path) -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _east(run, font="微软雅黑"):
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font)

    def h1(text: str):
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.name = "Calibri"
            _east(run)
        return p

    def para(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    # Title
    t = doc.add_heading("TikTok Shop 东南亚热门类目选品数据分析报告", level=0)
    for run in t.runs:
        run.font.name = "Calibri"
        _east(run)
    para(f"市场：泰国站（TH）   |   数据日期：{utc_today()}   |   数据源：SocialCrawl API（真实采集）")
    para("")

    # 1. Summary
    h1("一、摘要")
    total_sold = sum(s["total_sold"] for s in summaries)
    total_gmv = sum(s["total_gmv"] for s in summaries)
    top_cat = summaries[0]["category"] if summaries else "-"
    hottest = max(products, key=lambda p: p["sold_count"] or 0) if products else None
    para(f"本次共采集泰国站 {len(products)} 款在售商品，覆盖 12 个东南亚热门类目；"
         f"样本累计销量 {total_sold:,} 件，预估 GMV 约 {total_gmv:,.0f} 泰铢。")
    para(f"销量规模最大的类目是「{top_cat}」；单品销量冠军为「{hottest['title'][:50]}…」"
         f"（累计 {hottest['sold_count']:,} 件，现价 ฿{hottest['price']:.2f}）。")
    para("核心结论：")
    for s in summaries[:5]:
        para(f"  • {s['category']}：样本 {s['count']} 款，均价 ฿{s['avg_price']:.0f}，"
             f"总销量 {s['total_sold']:,} 件，平均评分 {s['avg_rating']}。", bold=False)

    # 2. Method
    h1("二、数据来源与方法")
    para("数据通过 SocialCrawl 的 TikTok Shop 搜索接口实时采集（region=TH），"
         "每个类目取搜索热度前 30 名商品；价格为泰铢（THB），销量为商品累计销量。")
    para("分析指标：折扣率（现价 vs 原价）、预估GMV（销量×现价）、蓝海分"
         "（40% 需求 + 30% 质量 + 30% 低竞争，仅在销量 100~30 万区间内计算）。")

    # 3. Category overview table
    h1("三、类目总览")
    table = doc.add_table(rows=1, cols=9)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["类目", "商品数", "均价(THB)", "中位价", "最高销量", "总销量", "平均评分", "折扣率", "预估GMV(THB)"]
    for i, htext in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = htext
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r in summaries:
        cells = table.add_row().cells
        values = [r["category"], str(r["count"]), f'{r["avg_price"]:.0f}', f'{r["median_price"]:.0f}',
                  f'{r["max_sold"]:,}', f'{r["total_sold"]:,}', f'{r["avg_rating"]}',
                  f'{r["avg_discount"]:.0%}', f'{r["total_gmv"]:,.0f}']
        for i, v in enumerate(values):
            cells[i].text = v

    # 4. Category insights
    h1("四、重点类目洞察")
    insights = {
        "美妆个护-护肤": "护肤是东南亚 TikTok Shop 体量最大的品类之一，精华、面霜类目出现大量"
                          "「买一送一」「组合装」打法，单品销量可达数十万件；高销量单品普遍采用"
                          "低价引流 + 高折后价的组合，适合达人带货矩阵。",
        "美妆个护-彩妆": "彩妆类目上新快、内容属性强，与美妆博主内容高度绑定；定价多集中在"
                          "中低价带（฿100-300），复购率高。",
        "服饰-女装": "女装以连衣裙、套装为主，价格带跨度大；高销量款多为宽松版型、多色可选，"
                      "视频展示（试穿对比）转化明显。",
        "服饰-男装": "男装以基础款 T 恤、衬衫为主，主打「免烫」「百搭」卖点；单品价格低、"
                      "走量属性强，适合低价跑量。",
        "3C数码-手机配件": "手机壳等配件标准化程度高、客单价低、内容展示直观，是典型的"
                            "内容电商跑量类目；爆款依赖外观设计与联名。",
        "家居生活-家居装饰": "家居装饰类目客单价相对较高，需展示真实使用场景；收纳、氛围灯等"
                              "小件适合短视频种草。",
        "食品饮料-零食": "零食类目复购率高、冲动消费强，直播试吃和吃播内容是核心转化手段；"
                          "注意食品资质与保质期管控。",
        "健康保健-保健品": "保健品类目客单价高、毛利高，但需要 FDA/注册资质背书；"
                            "「维C」「谷胱甘肽」等美白保健概念在泰国热度高。",
        "宠物用品": "东南亚养宠人群快速增长，猫砂、宠物食品消耗量大且复购稳定；"
                     "宠物用品的忠诚度强，适合建立店铺复购。",
        "运动户外-健身器材": "瑜伽垫等居家健身器材在泰国需求旺盛，头部单品累计销量超 17 万件；"
                               "定价带 ฿40-300 为主，适合捆绑配件提高客单价。",
        "时尚饰品-珠宝": "饰品客单价低、款式多，适合直播展示与「盲盒」「组合」玩法；"
                           "925 银、钛钢等材质概念是卖点。",
        "母婴用品": "母婴类目决策谨慎、复购强，家长更看重安全资质与口碑；"
                     "纸尿裤、婴儿护理等高消耗品适合订阅式复购。",
    }
    for s in summaries:
        para(f"{s['category']}：{insights.get(s['category'], '待补充')}", bold=True)

    # 5. Top products
    h1("五、热销商品 Top 15")
    top = sorted(products, key=lambda p: p["sold_count"] or 0, reverse=True)[:15]
    table2 = doc.add_table(rows=1, cols=7)
    table2.style = "Light Grid Accent 1"
    h2 = ["排名", "商品（节选）", "类目", "售价(THB)", "销量", "评分", "店铺"]
    for i, htext in enumerate(h2):
        cell = table2.rows[0].cells[i]
        cell.text = htext
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, p in enumerate(top, 1):
        cells = table2.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(p["title"])[:60]
        cells[2].text = p["category"]
        cells[3].text = f'{p["price"]:.2f}'
        cells[4].text = f'{p["sold_count"]:,}'
        cells[5].text = f'{p["rating"] or 0:.1f}'
        cells[6].text = str(p["seller_name"])[:25]

    # 6. Blue ocean picks
    h1("六、蓝海选品建议")
    shop_freq: dict = {}
    for p in products:
        shop_freq[p["seller_id"]] = shop_freq.get(p["seller_id"], 0) + 1
    max_freq = max(shop_freq.values()) if shop_freq else 1
    scored = []
    for p in products:
        if 100 <= (p["sold_count"] or 0) <= 300000:
            scored.append((blue_ocean_score(p, shop_freq, max_freq), p))
    scored.sort(key=lambda x: x[0], reverse=True)
    para("以下商品满足「有销量验证（≥100件）、价格有利润空间、竞争店铺少」的蓝海特征：")
    for i, (score, p) in enumerate(scored[:8], 1):
        para(f"  {i}. {str(p['title'])[:55]}｜{p['category']}｜฿{p['price']:.2f}｜"
             f"销量{p['sold_count']:,}｜评分{p['rating'] or 0:.1f}｜蓝海分{score:.3f}")

    # 7. Notes
    h1("七、风险与注意事项")
    for note in [
        "1. 销量为商品累计销量（非 7 日/30 日增量），只能反映规模，不能直接代表当前增速；",
        "2. 价格为泰铢（THB），汇率约 1 USD ≈ 33-35 THB，核算毛利时需按实时汇率换算；",
        "3. 高销量商品可能已进入价格战阶段，上新前需复核近 7 日销量走势；",
        "4. 保健品、母婴、食品类目需确认泰国 FDA 等合规资质后方可上架；",
        "5. 本报告基于单次截面采集（每类目 Top30），用于选品方向参考，建议结合多日趋势验证。",
    ]:
        para(note)

    doc.save(out)
    print(f"docx saved: {out}")



def build_html(products: list[dict], summaries: list[dict], out: Path) -> None:
    """Self-contained dashboard HTML (no external CDN), rendered from DB data."""
    import json

    shop_freq: dict = {}
    for p in products:
        shop_freq[p["seller_id"]] = shop_freq.get(p["seller_id"], 0) + 1
    max_freq = max(shop_freq.values()) if shop_freq else 1
    scored = []
    for p in products:
        if 100 <= (p["sold_count"] or 0) <= 300000:
            scored.append((blue_ocean_score(p, shop_freq, max_freq), p))
    scored.sort(key=lambda x: x[0], reverse=True)
    top = sorted(products, key=lambda p: p["sold_count"] or 0, reverse=True)[:10]

    insights = {
        "美妆个护-护肤": "精华、面霜类目普遍采用「买一送一」「组合装」打法，高销量单品常以低价引流+高折后价组合，适合达人带货矩阵。",
        "美妆个护-彩妆": "上新快、内容属性强，与美妆博主强绑定；定价集中在฿100-300 中低价带，复购率高。",
        "服饰-女装": "以连衣裙、套装为主，高销量款多为宽松版型、多色可选，试穿对比类视频转化明显。",
        "服饰-男装": "以基础款 T 恤、衬衫为主，主打「免烫」「百搭」，单价低、走量属性强。",
        "3C数码-手机配件": "标准化程度高、客单价低、内容展示直观，爆款依赖外观设计与联名。",
        "家居生活-家居装饰": "客单价相对较高，需展示真实使用场景；收纳、氛围灯等小件适合短视频种草。",
        "食品饮料-零食": "复购率高、冲动消费强，直播试吃与吃播内容是核心转化手段。",
        "健康保健-保健品": "客单价与毛利高，需 FDA/注册资质背书；「维C」「谷胱甘肽」等美白保健概念热度高。",
        "宠物用品": "养宠人群快速增长，猫砂、宠物食品消耗大且复购稳定，适合建立店铺复购。",
        "运动户外-健身器材": "瑜伽垫等居家健身器材需求旺盛，头部单品累计销量超 17 万件；可捆绑配件提高客单价。",
        "时尚饰品-珠宝": "客单价低、款式多，适合直播展示与「组合」「盲盒」玩法；注意低价引流款混杂。",
        "母婴用品": "决策谨慎、复购强，家长看重安全资质与口碑；纸尿裤、护理等高消耗品适合订阅式复购。",
    }

    payload = {
        "date": utc_today(),
        "count": len(products),
        "total_sold": sum(p["sold_count"] or 0 for p in products),
        "total_gmv": sum(p["gmv"] for p in products),
        "categories": summaries,
        "top": [
            {"rank": i, "title": str(p["title"])[:70], "category": p["category"],
             "price": round(p["price"] or 0, 2), "sold": p["sold_count"] or 0,
             "rating": p["rating"] or 0, "shop": str(p["seller_name"])[:30]}
            for i, p in enumerate(top, 1)
        ],
        "blue": [
            {"rank": i, "score": round(sc, 3), "title": str(p["title"])[:70],
             "category": p["category"], "price": round(p["price"] or 0, 2),
             "sold": p["sold_count"] or 0, "rating": p["rating"] or 0,
             "shop": str(p["seller_name"])[:30]}
            for i, (sc, p) in enumerate(scored[:10], 1)
        ],
        "insights": [{"category": s["category"],
                      "text": insights.get(s["category"], "")} for s in summaries],
    }

    html = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>TikTok Shop 东南亚热门类目选品数据分析</title>
<style>
  :root{--brand:#1f6feb;--bg:#f5f7fb;--card:#fff;--ink:#1f2328;--muted:#6b7280;--line:#e5e7eb;}
  *{box-sizing:border-box;margin:0;padding:0;}
  body{font-family:"Segoe UI","Microsoft YaHei",system-ui,sans-serif;background:var(--bg);color:var(--ink);line-height:1.6;}
  .wrap{max-width:1080px;margin:0 auto;padding:32px 20px 64px;}
  header h1{font-size:26px;letter-spacing:.5px;}
  .meta{color:var(--muted);font-size:13px;margin-top:6px;}
  .kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin:24px 0;}
  .kpi{background:var(--card);border:1px solid var(--line);border-radius:12px;padding:16px;}
  .kpi .num{font-size:24px;font-weight:700;color:var(--brand);font-variant-numeric:tabular-nums;}
  .kpi .lbl{font-size:12px;color:var(--muted);margin-top:2px;}
  .card{background:var(--card);border:1px solid var(--line);border-radius:12px;padding:20px;margin:18px 0;}
  h2{font-size:18px;margin-bottom:14px;}
  .bar-row{display:grid;grid-template-columns:150px 1fr 92px;gap:10px;align-items:center;margin:9px 0;font-size:13px;}
  .bar-track{background:#eef2f7;border-radius:6px;height:18px;overflow:hidden;}
  .bar-fill{height:100%;background:linear-gradient(90deg,#1f6feb,#54aeff);border-radius:6px;}
  table{width:100%;border-collapse:collapse;font-size:13px;}
  th,td{padding:8px 10px;border-bottom:1px solid var(--line);text-align:left;vertical-align:top;}
  th{background:#f8fafc;font-weight:600;white-space:nowrap;}
  td.num,th.num{text-align:right;font-variant-numeric:tabular-nums;white-space:nowrap;}
  .tag{display:inline-block;background:#e8f0fe;color:#1f6feb;border-radius:99px;padding:2px 10px;font-size:12px;white-space:nowrap;}
  .insight{margin:10px 0;padding:12px 14px;background:#f8fafc;border-left:3px solid var(--brand);border-radius:6px;font-size:13.5px;}
  .insight b{color:var(--brand);}
  footer{color:var(--muted);font-size:12px;margin-top:28px;border-top:1px solid var(--line);padding-top:14px;}
  @media(max-width:720px){.kpis{grid-template-columns:repeat(2,1fr);}.bar-row{grid-template-columns:110px 1fr 70px;}}
</style>
</head>
<body>
<div class="wrap">
  <header>
    <h1>TikTok Shop 东南亚热门类目选品数据分析</h1>
    <div class="meta" id="meta"></div>
  </header>
  <div class="kpis" id="kpis"></div>
  <div class="card"><h2>类目销量总览</h2><div id="cats"></div></div>
  <div class="card"><h2>热销商品 Top 10</h2><div id="top"></div></div>
  <div class="card"><h2>蓝海机会 Top 10</h2><div id="blue"></div></div>
  <div class="card"><h2>类目洞察</h2><div id="insights"></div></div>
  <footer>
    数据来源：SocialCrawl API（TikTok Shop 泰国站实时搜索，每类目 Top30）· 币种：泰铢（THB）·
    预估 GMV = 累计销量 × 现价 · 本页面由 tools/build_sea_report.py 自动生成，仅用于选品方向参考。
  </footer>
</div>
<script>
const DATA = __DATA__;
const fmt = (n) => n >= 1e8 ? (n/1e8).toFixed(2)+' 亿' : n >= 1e4 ? (n/1e4).toFixed(1)+' 万' : n.toLocaleString();
document.getElementById('meta').textContent = '市场：泰国站（TH） ｜ 数据日期：' + DATA.date +
  ' ｜ 数据源：SocialCrawl API 真实采集 ｜ 样本：' + DATA.count + ' 款商品';

const kpiBox = [
  [DATA.count, '采集商品（款）'],
  [fmt(DATA.total_sold), '样本累计销量（件）'],
  ['฿ ' + fmt(Math.round(DATA.total_gmv)), '预估 GMV（泰铢）'],
  [DATA.categories.length, '覆盖类目'],
];
document.getElementById('kpis').innerHTML = kpiBox.map(([num, lbl]) =>
  '<div class="kpi"><div class="num">' + num + '</div><div class="lbl">' + lbl + '</div></div>').join('');

const maxSold = Math.max(...DATA.categories.map(c => c.total_sold), 1);
document.getElementById('cats').innerHTML = DATA.categories.map(c =>
  '<div class="bar-row"><div>' + c.category + '</div>' +
  '<div class="bar-track"><div class="bar-fill" style="width:' + (c.total_sold / maxSold * 100).toFixed(1) + '%"></div></div>' +
  '<div class="num" style="text-align:right">' + fmt(c.total_sold) + '</div></div>').join('');

function tableHtml(rows, cols) {
  return '<table><thead><tr>' + cols.map(c => '<th class="' + (c.num ? 'num' : '') + '">' + c.label + '</th>').join('') +
    '</tr></thead><tbody>' + rows.map(r =>
    '<tr>' + cols.map(c => '<td class="' + (c.num ? 'num' : '') + '">' + r[c.key] + '</td>').join('') +
    '</tr>').join('') + '</tbody></table>';
}

document.getElementById('top').innerHTML = tableHtml(
  DATA.top.map(r => ({
    rank: r.rank, title: r.title, category: '<span class="tag">' + r.category + '</span>',
    price: '฿' + r.price.toLocaleString(), sold: fmt(r.sold), rating: r.rating.toFixed(1), shop: r.shop
  })),
  [{key:'rank',label:'#'},{key:'title',label:'商品'},{key:'category',label:'类目'},
   {key:'price',label:'售价',num:true},{key:'sold',label:'累计销量',num:true},
   {key:'rating',label:'评分',num:true},{key:'shop',label:'店铺'}]
);

document.getElementById('blue').innerHTML = tableHtml(
  DATA.blue.map(r => ({
    rank: r.rank, score: r.score.toFixed(3), title: r.title,
    category: '<span class="tag">' + r.category + '</span>',
    price: '฿' + r.price.toLocaleString(), sold: fmt(r.sold), rating: r.rating.toFixed(1), shop: r.shop
  })),
  [{key:'rank',label:'#'},{key:'score',label:'蓝海分',num:true},{key:'title',label:'商品'},
   {key:'category',label:'类目'},{key:'price',label:'售价',num:true},{key:'sold',label:'销量',num:true},
   {key:'rating',label:'评分',num:true},{key:'shop',label:'店铺'}]
);

document.getElementById('insights').innerHTML = DATA.insights.map(i =>
  '<div class="insight"><b>' + i.category + '</b>：' + i.text + '</div>').join('');
</script>
</body>
</html>"""

    html = html.replace("__DATA__", json.dumps(payload, ensure_ascii=False))
    out.write_text(html, encoding="utf-8")
    print(f"html saved: {out}")

def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--db", default="data/tiktok_shop.db")
    ap.add_argument("--outdir", default="reports")
    args = ap.parse_args()

    products = load_products(args.db)
    if not products:
        print("no products found for SEA categories; run tools/collect_socialcrawl.py first")
        return 1
    summaries = category_summary(products)
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    day = utc_today()
    build_xlsx(products, summaries, outdir / f"东南亚热门类目数据分析_{day}.xlsx")
    build_docx(products, summaries, outdir / f"TikTokShop东南亚热门类目选品分析报告_{day}.docx")
    build_html(products, summaries, outdir / "sea_report.html")
    return 0


if __name__ == "__main__":
    sys.exit(main())
